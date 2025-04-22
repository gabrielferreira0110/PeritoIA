from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime

def generate_docx_report(title, sections_data, author_name):
    """
    Gera um documento DOCX com base nas seções do laudo pericial
    
    Args:
        title (str): Título do laudo
        sections_data (dict): Dicionário com os dados das seções
        author_name (str): Nome do autor/perito
        
    Returns:
        BytesIO: Buffer contendo o documento DOCX
    """
    doc = Document()
    
    # Configurar margens do documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Adicionar cabeçalho
    header = doc.add_heading(title.upper(), level=0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adicionar metadados
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Perito: {author_name}").bold = True
    p.add_run(f"\nData: {datetime.now().strftime('%d/%m/%Y')}")
    
    # Adicionar separador
    doc.add_paragraph("__" * 30)
    
    # Mapeamento de nomes de seções para títulos formatados
    section_titles = {
        'preambulo': 'PREÂMBULO',
        'palavras_chave': 'PALAVRAS-CHAVE',
        'apresentacao_demanda': 'APRESENTAÇÃO DA DEMANDA',
        'objeto_pericia': 'OBJETO DA PERÍCIA',
        'metodologia': 'METODOLOGIA',
        'descricao': 'DESCRIÇÃO',
        'discussao': 'DISCUSSÃO',
        'conclusao': 'CONCLUSÃO'
    }
    
    # Ordem das seções no documento
    section_order = [
        'preambulo', 
        'palavras_chave', 
        'apresentacao_demanda',
        'objeto_pericia',
        'metodologia',
        'descricao',
        'discussao',
        'conclusao'
    ]
    
    # Adicionar cada seção ao documento
    for section_key in section_order:
        if section_key in sections_data and sections_data[section_key]:
            # Adicionar título da seção
            section_title = section_titles.get(section_key, section_key.upper())
            heading = doc.add_heading(section_title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Adicionar conteúdo da seção
            content = sections_data[section_key]
            p = doc.add_paragraph(content)
            
            # Adicionar quebra de página após cada seção (exceto a última)
            if section_key != section_order[-1]:
                doc.add_page_break()
    
    # Adicionar rodapé
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento gerado pelo sistema Perito IA")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Salvar o documento em um buffer de memória
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

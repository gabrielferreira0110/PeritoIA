import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def generate_report_docx(report, sections):
    """
    Generate a DOCX file from a report and its sections
    
    Args:
        report (Report): The report object
        sections (list): List of ReportSection objects
    
    Returns:
        str: Path to the generated DOCX file
    """
    try:
        # Create a new document
        doc = Document()
        
        # Set document properties
        core_properties = doc.core_properties
        core_properties.title = report.title
        core_properties.author = "Perito IA"
        core_properties.language = "pt-BR"
        
        # Configure document style
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)
        
        # Add title
        title = doc.add_heading(report.title.upper(), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sort sections in the correct order
        section_order = {
            'preambulo': 1,
            'palavras_chave': 2,
            'apresentacao_demanda': 3,
            'objeto_pericia': 4,
            'metodologia': 5,
            'descricao': 6,
            'discussao': 7,
            'conclusao': 8
        }
        
        sorted_sections = sorted(
            sections, 
            key=lambda s: section_order.get(s.section_type, 99)
        )
        
        # Section headers in Portuguese
        section_headers = {
            'preambulo': 'PREÂMBULO',
            'palavras_chave': 'PALAVRAS-CHAVE',
            'apresentacao_demanda': 'APRESENTAÇÃO DA DEMANDA',
            'objeto_pericia': 'OBJETO DA PERÍCIA',
            'metodologia': 'METODOLOGIA',
            'descricao': 'DESCRIÇÃO',
            'discussao': 'DISCUSSÃO',
            'conclusao': 'CONCLUSÃO'
        }
        
        # Add each section to the document
        for section in sorted_sections:
            # Add section header
            header_text = section_headers.get(section.section_type, section.section_type.upper())
            heading = doc.add_heading(header_text, level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add section content
            paragraphs = section.content.split("\n")
            for p in paragraphs:
                if p.strip():
                    para = doc.add_paragraph(p.strip())
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Add date and signature at the end
        doc.add_paragraph('')  # Add some space
        date_text = f"{datetime.now().strftime('%d de %B de %Y')}"
        date_para = doc.add_paragraph(date_text)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('')  # Add some space
        signature_para = doc.add_paragraph("_______________________________")
        signature_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        name_para = doc.add_paragraph("Perito Judicial")
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create directory for reports if it doesn't exist
        reports_dir = os.path.join(os.getcwd(), "uploads", "reports")
        os.makedirs(reports_dir, exist_ok=True)
        
        # Save the document
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        docx_path = os.path.join(reports_dir, f"laudo_pericial_{report.id}_{timestamp}.docx")
        doc.save(docx_path)
        
        return docx_path
    except Exception as e:
        logging.error(f"Error generating DOCX: {str(e)}")
        raise
